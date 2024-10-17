import os, pathlib, sys
current_dir = os.getcwd()
dir_to_add = pathlib.Path(current_dir).parent.as_posix()
sys.path.append(dir_to_add)
#---------------------------------------
from docx import Document
from dateparser import parse
import httpx
import config
from country_named_entity_recognition import find_countries
from iteration_utilities import unique_everseen
from countryguess import guess_country
from dateparser import parse
from datetime import timedelta
import bot_static


tag_dict = {
    'economic_586': [x.strip() for x in bot_static.economic.split(',')],
    'monetary policy_587': [x.strip() for x in bot_static.monetary_policy.split(',')],
    'political_588': [x.strip() for x in bot_static.political.split(',')]
}


def get_tags(text):
    tags = []
    for x in tag_dict:
        for tag_name in tag_dict[x]:
            if tag_name.lower() in text.lower():
                tags.append(x.split('_')[-1])
    tags = list(set(tags))
    return tags


def get_categories():
    print(f'Getting Categories')
    categories = []
    num = 1
    while True:
        print(f'Page: {num}')
        response = httpx.get(
        f'{config.wordpress_host}/wp-json/wp/v2/categories?per_page=100&page={num}', 
        auth=(config.wordress_pswd, config.wordress_pswd))
        json_data = response.json()
        if len(json_data) > 0:
           [categories.append(x) for x in json_data]
        else:
            break
        num += 1
    return categories


def get_country_id(c_name, categories: list):
    c_name = 'The Philippines' if c_name == 'Philippines' else c_name
    c_name = 'The Dominican Republic' if c_name == 'Dominican Republic' else c_name
    c_name = 'Cape Verde' if c_name == 'Cabo Verde' else c_name
    c_name = 'Ivory Coast' if c_name == "Cote d'Ivoire" else c_name
    c_name = 'Turkey' if c_name == 'Türkiye' else c_name
    c_name = 'Bosnia' if c_name == 'Bosnia and Herzegovina' else c_name
    c_name = 'Kyrgyzstan' if c_name == 'Kyrgyz Republic' else c_name
    for x in categories:
        if x['name'].lower() in c_name.lower():
            return str(x['id'])


def save_data(country, title, text, date, all_posts: list):
    all_posts.append({
        'categories': country,
        'title': title,
        'content': text,
        'date': date,
        'tags': get_tags(text)
    })


def parse_text_into_countries(text, date, all_posts: list):
    places = find_countries(text, is_ignore_case=True)
    countries = []
    if len(places) > 0:
        for place in places:
            try:
                country = place[0].common_name
                countries.append(country)
            except:
                country = guess_country(place[0].name, attribute='name_short')
                countries.append(country)
    else:
        country = guess_country(text, attribute='name_short')
        if country:
            countries.append(country)
    countries = list(set(countries))
    for x in countries:
        if x in bot_static.country_to_ignore:
            countries.remove(x)
    if len(countries) > 0:
        save_data(countries, countries[0], text, date, all_posts)
  
  
def update_date(date: str, index_number: int) -> str:
    #get the date based on the current index number and the starting date
    date_str = (timedelta(index_number) + parse(date))
    if date_str.isoweekday() == 6: #if saturday change to friday
        print('it is saturday') 
        index_number += 2
        date_str = (timedelta(index_number) + parse(date))
    elif date_str.isoweekday() == 7: #if sunday change to friday
        print('it is sunday') 
        index_number += 1
        date_str = (timedelta(index_number) + parse(date))
    index_number += 1 #add one for the next day (this will be applied on the next day)
    date_str = date_str.strftime('%Y-%m-%dT%H:%M:%S')
    
    return date_str, index_number


def to_ignore_parargraph(paragraph_value, to_ignores: list[str]):
    #tg -> to_ignore, tgs -> to_ignores
    x_value = paragraph_value.text
    check_if_to_ignore = lambda tg: \
    (tg.lower() in x_value.lower()) or (len(x_value) < 10) or (x_value.isspace())
    results = list(filter(check_if_to_ignore, to_ignores))
    if results:
        return True
    else:
        return False


def bulk_parse_document(start_date: str, f_content):
    print(f'Parsing document')
    all_posts: list[dict] = []
    new_day_text = "Good morning – Time for the daily market update"
    to_ignores: list[str] = ["IN CIS:", "IN ASIA:", "IN MEA:", "IN LATAM:", "Good morning – Time for the daily market update"]
    index_number: int = 0
    
    content = Document(f_content)
    paragraphs = content.paragraphs
    for paragraph_value in paragraphs:
        if (new_day_text.lower() in paragraph_value.text.lower()): #New data day shows up
            date, index_number = update_date(start_date, index_number) #Get new date
        if not to_ignore_parargraph(paragraph_value, to_ignores): #Filter the values
            parse_text_into_countries(paragraph_value.text, date, all_posts) #And do some parsing
    all_posts = list(unique_everseen(all_posts))
    return all_posts


def parse_document(start_date: str, content):
    print(f'Parsing document')
    all_posts = []
    to_ignores = ["IN CIS:", "IN ASIA:", "IN MEA:", "IN LATAM:", "Good morning – Time for the daily market update"]
    document = Document(content)
    date_str = parse(start_date) if start_date else parse('today')
    date_str = date_str.strftime('%Y-%m-%dT%H:%M:%S')
    for x in document.paragraphs:
        to_ignore = False
        for tg in to_ignores:
            if tg.lower() in x.text.lower():
                to_ignore = True
                break
        if not to_ignore:
            parse_text_into_countries(x.text, date_str, all_posts)
    all_posts = list(unique_everseen(all_posts))
    return all_posts


def send_to_wordpress(post_dict: list, categories_list):
    categories = [
        get_country_id(cat, categories_list) for cat in post_dict.get('categories')]
    while None in categories:
        categories.remove(None)
    if len(categories) > 0:
        json_data = {
            'title': post_dict.get('title'),
            'date': post_dict.get('date'),
            'content': post_dict.get('content'),
            'categories': categories,
            'status': 'publish',
            'tags': post_dict.get('tags')
        }
        response = httpx.post(
        f'{config.wordpress_host}/wp-json/wp/v2/posts', 
        auth=(config.wordress_username, config.wordress_pswd), json=json_data)
        print(response)
        if response.status_code == 201:
            print("Post sent")
        else:
            print(json_data['categories'])


def engine():
    categories = get_categories()
    with open('/home/khaliq/Desktop/DMU Sep 17-19.docx', 'rb') as f:
        #content = Document(f)
        all_posts = bulk_parse_document('September 17th 2024', f)
        # for post in all_posts[:3]:
        #     send_to_wordpress(post, categories)
        return all_posts
        
    # if len(all_posts) > 0:
    #     send_to_wordpress(all_posts)
        
        
if __name__ == '__main__':
    engine()